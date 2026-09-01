"""PDF, DOCX, and image text extraction."""

from __future__ import annotations

from pathlib import Path
from typing import Any


def extractPdf(readPath: Path) -> str:
    import pdfplumber

    readParts = []
    with pdfplumber.open(str(readPath)) as readPdf:
        for readPage in readPdf.pages:
            readText = (readPage.extract_text() or "").strip()
            if readText:
                readParts.append(readText)
            for readTable in readPage.extract_tables() or []:
                for readRow in readTable or []:
                    readCells = [str(readCell).strip() if readCell is not None else "" for readCell in readRow]
                    if any(readCells):
                        readParts.append(" | ".join(readCells))
    return "\n\n".join(readParts).strip()


def extractDocx(readPath: Path) -> str:
    from docx import Document

    readDoc = Document(str(readPath))
    readParts = [readParagraph.text.strip() for readParagraph in readDoc.paragraphs
                 if readParagraph.text.strip()]
    for readTable in readDoc.tables:
        for readRow in readTable.rows:
            readCells = [readCell.text.strip() for readCell in readRow.cells]
            if any(readCells):
                readParts.append(" | ".join(readCells))
    return "\n".join(readParts)


def readOcr(readImage: Any) -> str:
    try:
        from paddleocr import PaddleOCR
    except ImportError as readError:
        raise RuntimeError("PaddleOCR is not installed") from readError
    createOcr = PaddleOCR(lang="ch", use_doc_orientation_classify=True, use_doc_unwarping=True)
    readResult = createOcr.predict(readImage) if hasattr(createOcr, "predict") else createOcr.ocr(readImage)
    readLines = []
    for readItem in readResult or []:
        readData = getattr(readItem, "res", None)
        if not isinstance(readData, dict) and hasattr(readItem, "json"):
            readData = readItem.json().get("res", {})
        if isinstance(readData, dict):
            readLines.extend(str(readText).strip() for readText in readData.get("rec_texts", []) if readText)
        elif isinstance(readItem, list):
            for readLine in readItem:
                if isinstance(readLine, (list, tuple)) and len(readLine) > 1:
                    readValue = readLine[1][0] if isinstance(readLine[1], (list, tuple)) else readLine[1]
                    if readValue:
                        readLines.append(str(readValue).strip())
    return "\n".join(readLines)


def extractImage(readPath: Path) -> str:
    import numpy
    from PIL import Image

    readImage = Image.open(readPath).convert("RGB")
    readImage.thumbnail((2000, 2000), Image.Resampling.LANCZOS)
    return readOcr(numpy.array(readImage, dtype=numpy.uint8))


def extractText(readPath: Path) -> str:
    readSuffix = readPath.suffix.lower()
    if readSuffix == ".pdf":
        return extractPdf(readPath)
    if readSuffix == ".docx":
        return extractDocx(readPath)
    if readSuffix in {".png", ".jpg", ".jpeg"}:
        return extractImage(readPath)
    raise ValueError(f"unsupported resume format: {readSuffix or '(none)'}")
